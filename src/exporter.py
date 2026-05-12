"""导出模块 - DOCX 输出完整文章段落"""
import os
from subtitle_downloader import SubtitleEntry
from translator import subtitles_to_paragraphs, _clean_chinese_spacing, _clean_chinese_punctuation


def export_docx(entries, output_path, title="YouTube 字幕", subtitle_type=""):
    """导出为 DOCX 格式（排版精美）"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.8
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 标题
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 字幕类型说明
    if subtitle_type:
        type_para = doc.add_paragraph()
        type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = type_para.add_run(subtitle_type)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph("")

    # 按段落输出
    raw_paragraphs = subtitles_to_paragraphs(entries)
    paragraphs = raw_paragraphs.split("\n\n")

    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        para_text = _clean_chinese_punctuation(para_text)
        para_text = _clean_chinese_spacing(para_text)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(22)
        run = p.add_run(para_text)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.save(output_path)


def export_docx_text(text, output_path, title="YouTube 字幕", subtitle_type=""):
    """直接从完整文本导出 DOCX（用于 Whisper 转录结果）"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.8
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    if subtitle_type:
        type_para = doc.add_paragraph()
        type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = type_para.add_run(subtitle_type)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph("")

    # 处理文本
    text = _clean_chinese_punctuation(text)
    text = _clean_chinese_spacing(text)

    # 按段落分割
    paragraphs = text.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(22)
        run = p.add_run(para_text)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.save(output_path)


def export_all(entries, output_dir, base_name, subtitle_type=""):
    """导出 DOCX"""
    os.makedirs(output_dir, exist_ok=True)
    generated = []
    path = os.path.join(output_dir, f"{base_name}.docx")
    export_docx(entries, path, title=base_name, subtitle_type=subtitle_type)
    generated.append(path)
    return generated
