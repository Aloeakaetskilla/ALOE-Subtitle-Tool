"""生成使用说明 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def create_guide(output_path):
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ──────────── 标题 ────────────
    title = doc.add_heading('ALOE 字幕提取软件 — 使用说明', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph('')

    # ──────────── 软件简介 ────────────
    doc.add_heading('一、软件简介', level=1)
    doc.add_paragraph(
        'ALOE 字幕提取软件是一款基于 Whisper 语音识别技术的视频字幕提取工具。'
        '支持从 YouTube 等视频平台下载音频，自动识别语音内容并翻译为中文，'
        '最终导出为 Word 文档格式。'
    )

    # ──────────── 功能特点 ────────────
    doc.add_heading('二、功能特点', level=1)
    features = [
        '支持 11 种主流语言的语音识别（自动检测 / 手动指定）',
        '识别语言包括：中文、英语、日语、韩语、法语、德语、西班牙语、俄语、葡萄牙语、阿拉伯语',
        '智能翻译系统，支持多引擎自动切换',
        '导出精美排版的 Word 文档',
        '支持单语 / 双语输出',
        '深色主题界面，操作简单直观',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    # ──────────── 系统要求 ────────────
    doc.add_heading('三、系统要求', level=1)
    reqs = [
        '操作系统：Windows 10 / 11（64 位）',
        '磁盘空间：至少 2 GB 可用空间（用于存放识别模型和临时文件）',
        '网络连接：需要联网（下载音频和翻译功能需要网络）',
    ]
    for r in reqs:
        doc.add_paragraph(r, style='List Bullet')

    # ──────────── 使用步骤 ────────────
    doc.add_heading('四、使用步骤', level=1)

    doc.add_heading('第 1 步：启动软件', level=2)
    doc.add_paragraph('双击「ALOE字幕提取软件.exe」即可启动，无需安装。')

    doc.add_heading('第 2 步：输入视频链接', level=2)
    doc.add_paragraph(
        '在「视频链接」输入框中粘贴 YouTube 视频链接，也可以点击右侧「粘贴」按钮从剪贴板获取。'
    )

    doc.add_heading('第 3 步：选择视频语言', level=2)
    doc.add_paragraph(
        '在「视频语言」下拉框中选择视频的原始语言。默认为「自动检测」，'
        '软件会自动判断视频语言。如果自动检测不准确，可以手动指定。'
    )

    doc.add_heading('第 3.5 步：选择识别精度', level=2)
    doc.add_paragraph(
        '在「识别精度」下拉框中选择模型大小。模型越大，识别越准确但速度越慢：'
    )
    model_opts = [
        'tiny — 极速，精度较低，适合快速预览',
        'base — 快速，精度一般',
        'small — 均衡，推荐大多数场景使用',
        'medium — 高精度，速度较慢',
        'large-v3 — 最高精度，速度很慢，适合重要内容',
    ]
    for m in model_opts:
        doc.add_paragraph(m, style='List Bullet')
    doc.add_paragraph('注意：首次使用某个模型时需要下载，之后会自动缓存。')

    doc.add_heading('第 4 步：选择输出语言', level=2)
    options = [
        '仅原文 — 只导出识别到的原始语言文字',
        '仅中文翻译 — 将识别结果翻译为中文后导出',
        '双语（原文+中文）— 同时导出原文和中文翻译',
    ]
    for o in options:
        doc.add_paragraph(o, style='List Bullet')

    doc.add_heading('第 5 步：设置保存位置', level=2)
    doc.add_paragraph(
        '默认保存到桌面的「YouTubeSubtitles」文件夹，也可以点击「浏览」自定义保存路径。'
    )

    doc.add_heading('第 6 步：开始提取', level=2)
    doc.add_paragraph(
        '点击「开始提取」按钮，软件将自动完成以下流程：'
    )
    steps = [
        '获取视频信息',
        '下载音频',
        '语音识别（首次使用需下载识别模型，请耐心等待）',
        '翻译（如选择了中文输出）',
        '生成 Word 文档',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {s}')

    doc.add_paragraph('处理完成后，会提示是否打开输出目录查看生成的文件。')

    # ──────────── 常见问题 ────────────
    doc.add_heading('五、常见问题', level=1)

    doc.add_heading('Q: 提示「未找到 yt-dlp」怎么办？', level=3)
    doc.add_paragraph(
        '请确保 yt-dlp.exe 与软件放在同一目录下。如果没有该文件，'
        '可以从 https://github.com/yt-dlp/yt-dlp/releases 下载最新版本。'
    )

    doc.add_heading('Q: 语音识别很慢怎么办？', level=3)
    doc.add_paragraph(
        'Whisper 语音识别使用 CPU 运行，速度取决于电脑性能。'
        '较长的视频可能需要较长时间。首次使用时需要下载模型，之后会自动缓存。'
    )

    doc.add_heading('Q: 识别结果不准确怎么办？', level=3)
    doc.add_paragraph(
        '可以尝试手动指定视频语言（而非自动检测），这通常能提高识别准确率。'
        '另外，音质较差的视频识别效果也会受影响。'
    )

    doc.add_heading('Q: 翻译失败了怎么办？', level=3)
    doc.add_paragraph(
        '软件内置多个翻译引擎，会自动切换。如果翻译失败，原文仍会正常导出。'
        '请检查网络连接是否正常。'
    )

    # ──────────── 输出文件说明 ────────────
    doc.add_heading('六、输出文件说明', level=1)
    doc.add_paragraph('生成的文件为 .docx 格式（Word 文档），包含：')
    out_desc = [
        '视频标题作为文档标题',
        '原文或翻译内容（按段落排版）',
        '首行缩进、1.5 倍行距的精美排版',
    ]
    for d in out_desc:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('')

    # 底部
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('ALOE 字幕提取软件')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.save(output_path)
    print(f'使用说明已生成: {output_path}')


if __name__ == '__main__':
    import os
    out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), '使用说明.docx')
    create_guide(out)
